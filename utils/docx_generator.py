from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def generate_docx_report(content, output_path):
    """텍스트를 1분 스피치 DOCX 형식으로 변환"""
    doc = Document()
    
    # 제목 추가 (중앙 정렬)
    title = doc.add_heading('1분 스피치', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 부제목
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('주간 회의 발표 자료')
    subtitle_run.font.size = Inches(0.15)
    
    # 구분선
    doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 본문 내용 추가
    lines = content.split('\n')
    for line in lines:
        if line.strip():
            if line.startswith('#') or '제목' in line or '스피치' in line:
                # 헤딩으로 처리
                level = min(line.count('#'), 2) + 1
                heading_text = line.lstrip('#').strip()
                heading = doc.add_heading(heading_text, level)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                # 일반 텍스트로 처리
                para = doc.add_paragraph(line)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # 중요한 문장 강조
                if any(word in line for word in ['결론', '요약', '주요', '감사']):
                    for run in para.runs:
                        run.font.bold = True
    
    # 페이지 하단에 날짜 추가
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    from datetime import datetime
    footer_run = footer.add_run(f"작성일: {datetime.now().strftime('%Y년 %m월 %d일')}")
    footer_run.font.size = Inches(0.1)
    
    doc.save(output_path)
    print(f"DOCX 보고서 생성 완료: {output_path}")